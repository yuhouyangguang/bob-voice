import io
import json
import zipfile
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def document_basename(task):
    from ..utils import slug_filename

    meeting_date = (
        task.meeting.meeting_at.date()
        if task.meeting.meeting_at
        else task.created_at.date()
    )
    return f"{meeting_date:%Y%m%d}_{slug_filename(task.meeting.topic)}_发言整理"


def render_markdown(task):
    meeting = task.meeting
    when = meeting.meeting_at.strftime("%Y-%m-%d %H:%M") if meeting.meeting_at else "未填写"
    speakers = "、".join(
        item.get("name", "") if isinstance(item, dict) else str(item)
        for item in (meeting.key_speakers or [])
    ) or "未指定"
    lines = [
        f"# {when[:10]} {meeting.topic} - {speakers} 发言整理",
        "",
        "## 会议基本信息",
        "| 项目 | 内容 |",
        "|---|---|",
        f"| 会议形式 | {meeting.meeting_type} |",
        f"| 时间 | {when} |",
        f"| 地点 | {meeting.location or '未填写'} |",
        "",
        "## 校准说明",
        "> 本文由 BOB Voice 自动转写并经过术语规则校对，请结合原始录音复核。",
        "",
        "## 发言内容",
    ]
    for segment in meeting.segments:
        text = segment.corrected_text or segment.raw_text
        lines.extend(
            [
                "",
                f"**{segment.speaker_label}**（{format_timestamp(segment.start_time)}）：{text}",
            ]
        )
    lines.extend(["", "## 保密标注", "> 内部资料，禁止外传", ""])
    return "\n".join(lines)


def render_word(task):
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(task.meeting.topic + " 发言整理")
    set_run_font(title_run, "黑体", 22, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    when = (
        task.meeting.meeting_at.strftime("%Y年%m月%d日 %H:%M")
        if task.meeting.meeting_at
        else "时间未填写"
    )
    set_run_font(subtitle.add_run(when), "仿宋_GB2312", 12)

    heading = document.add_paragraph()
    set_run_font(heading.add_run("一、会议基本信息"), "黑体", 16)
    for label, value in (
        ("会议形式", task.meeting.meeting_type),
        ("会议地点", task.meeting.location or "未填写"),
    ):
        paragraph = document.add_paragraph()
        set_run_font(paragraph.add_run(f"{label}：{value}"), "仿宋_GB2312", 14)

    heading = document.add_paragraph()
    set_run_font(heading.add_run("二、发言内容"), "黑体", 16)
    for segment in task.meeting.segments:
        paragraph = document.add_paragraph()
        speaker_run = paragraph.add_run(f"{segment.speaker_label}：")
        set_run_font(speaker_run, "仿宋_GB2312", 14, bold=True)
        set_run_font(
            paragraph.add_run(segment.corrected_text or segment.raw_text),
            "仿宋_GB2312",
            14,
        )

    confidentiality = document.add_paragraph()
    confidentiality.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = confidentiality.add_run("★ 内部资料，禁止外传 ★")
    set_run_font(run, "黑体", 16, bold=True)
    run.font.color.rgb = RGBColor(196, 18, 48)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def set_run_font(run, name, size, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def render_json(task):
    payload = task.to_dict()
    payload["segments"] = [segment.to_dict() for segment in task.meeting.segments]
    if task.meeting.supervision:
        payload["supervision"] = task.meeting.supervision.to_dict()
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def render_zip(task):
    output = io.BytesIO()
    basename = document_basename(task)
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(f"{basename}.md", render_markdown(task).encode("utf-8"))
        archive.writestr(f"{basename}.docx", render_word(task).getvalue())
        archive.writestr(f"{basename}.json", render_json(task))
    output.seek(0)
    return output


def format_timestamp(seconds):
    seconds = max(int(seconds or 0), 0)
    hours, remainder = divmod(seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
